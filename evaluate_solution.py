import sys
import re
import os

try:
    import pypdf
except ImportError:
    pypdf = None
try:
    import docx
except ImportError:
    docx = None

def extract_text(file_path):
    if file_path.endswith('.pdf') and pypdf:
        reader = pypdf.PdfReader(file_path)
        return ' '.join([page.extract_text() for page in reader.pages])
    elif file_path.endswith('.docx') and docx:
        doc = docx.Document(file_path)
        text = ' '.join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += ' ' + ' '.join([cell.text for cell in row.cells])
        return text
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

def check_context_integrity(question_text, solution_text):
    # Extract Capitalized words that might be entities (length > 3 to avoid It, The, etc.)
    q_entities = set(re.findall(r'\b[A-Z][a-z]{3,}\b', question_text))
    s_entities = set(re.findall(r'\b[A-Z][a-z]{3,}\b', solution_text))
    
    # Common words to ignore in accounting contexts
    ignore = {"The", "This", "That", "Required", "Question", "Total", "Information", "March", "April", "September", "Statement", "Profit", "Loss", "Financial", "Position", "Account", "Accounts", "Assets", "Liabilities", "Equity", "Sales", "Cost", "Purchases", "Inventory", "Ordinary", "Current", "Group", "Retained", "Plant", "Equipment", "Value", "Fair", "Marks", "Note", "Working", "Workings", "Calculation"}
    q_entities = q_entities - ignore
    
    overlap = q_entities.intersection(s_entities)
    
    threshold = 3 # Expect at least 3 shared major entities
    if len(q_entities) < 3:
        threshold = len(q_entities)
        
    print(f"[SEARCH] Context Check: Found {len(overlap)} matching entities between question and solution.")
    if len(overlap) >= threshold:
        print(f"[PASS] Context Integrity Passed. Shared Entities: {list(overlap)[:5]}...")
        return True
    else:
        print(f"[FAIL] Context Integrity FAILED. Very low entity overlap.")
        print(f"     Could this solution be for a different question?")
        return False

def check_mathematical_integrity(solution_text):
    # Regex looks for the term, optionally some text/dots/spaces, and then a number/comma
    assets_match = re.search(r'Total\s+Assets.*?(?:GH¢|\s|:)+([\d,\.]+)', solution_text, re.IGNORECASE)
    equity_liab_match = re.search(r'Total\s+Equity\s+(?:and\s+)?Liabilities.*?(?:GH¢|\s|:)+([\d,\.]+)', solution_text, re.IGNORECASE)
    
    if not assets_match or not equity_liab_match:
        print(f"[SKIP] Math Check Skipped: Could not find strict 'Total Assets' or 'Total Equity and Liabilities' lines.")
        return True # Not all solutions have a balance sheet, so we skip rather than fail
        
    assets_val = float(assets_match.group(1).replace(',', ''))
    equity_val = float(equity_liab_match.group(1).replace(',', ''))
    
    print(f"[SEARCH] Math Check: Comparing Total Assets ({assets_val}) against Total Equity/Liab ({equity_val}).")
    
    if assets_val == equity_val:
        print(f"[PASS] Mathematical Integrity Passed. Balance Sheet balances perfectly.")
        return True
    else:
        print(f"[FAIL] Mathematical Integrity FAILED. Balance sheet does not balance!")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python evaluate_solution.py <question_file.html> <solution_file.txt/pdf/docx>")
        sys.exit(1)
        
    question_path = sys.argv[1]
    solution_path = sys.argv[2]
    
    if not os.path.exists(question_path):
        print(f"Error: Question file not found -> {question_path}")
        sys.exit(1)
    if not os.path.exists(solution_path):
        print(f"Error: Solution file not found -> {solution_path}")
        sys.exit(1)
    
    print(f"--- Starting Evaluation ---")
    print(f"Question: {os.path.basename(question_path)}")
    print(f"Solution: {os.path.basename(solution_path)}\n")
    
    q_text = extract_text(question_path)
    s_text = extract_text(solution_path)
    
    context_passed = check_context_integrity(q_text, s_text)
    math_passed = check_mathematical_integrity(s_text)
    
    print("\n--- Evaluation Result ---")
    if context_passed and math_passed:
        print("[SUCCESS] ALL CHECKS PASSED. Solution is safe to append.")
        sys.exit(0)
    else:
        print("[ERROR] CHECKS FAILED. Do not append this solution.")
        sys.exit(1)
