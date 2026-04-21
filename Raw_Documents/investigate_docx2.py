import docx
import sys

# Force utf-8 for stdout
sys.stdout.reconfigure(encoding='utf-8')

doc_path = r'C:\Users\avorx\OneDrive\Desktop\philipexamprep\Raw_Documents\ACC402D_SectionA_Solutions_v2.docx'
doc = docx.Document(doc_path)

print("Total tables:", len(doc.tables))
for i, table in enumerate(doc.tables):
    print(f"Table {i}:")
    for row in table.rows:
        print([cell.text for cell in row.cells])
