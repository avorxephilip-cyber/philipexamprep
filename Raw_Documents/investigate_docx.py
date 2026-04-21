import docx

doc_path = r'C:\Users\avorx\OneDrive\Desktop\philipexamprep\Raw_Documents\ACC402D_SectionA_Solutions_v2.docx'
doc = docx.Document(doc_path)

for i, para in enumerate(doc.paragraphs):
    if 'THE QUESTION' in para.text:
        print(f"--- Q block start at para {i} ---")
        for j in range(i, min(i+10, len(doc.paragraphs))):
            print(f"P{j}: {doc.paragraphs[j].text}")
        print("--- Tables nearby ---")
        # Find tables before the next question
        for table in doc.tables:
            # Just print all tables for now to see if they contain options
            pass
        break

print("Total tables:", len(doc.tables))
for i, table in enumerate(doc.tables[:3]):
    print(f"Table {i}:")
    for row in table.rows:
        print([cell.text for cell in row.cells])
